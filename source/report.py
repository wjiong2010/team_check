from docx import Document
import datetime, os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from kpi import season_date

replace_list = {
    '<<Name>>': '王炯',
    '<<Year>>': '2024',
    '<<Season>>': 'Q4',
    '<<S_Month>>': '222',
    '<<S_Day>>': '',
    '<<E_Month>>': '',
    '<<E_Day>>': '',
    '<<P_Score>>': '',
    '<<M_Score>>': '',
    '<<Key_Score>>': '0',
    '<<Total_Score>>': '',
    '<<Rank>>': '',
    '<<Level>>': '',
    '<<Comment>>': '',
    '<<Opinion>>': ''
}

season_cvt = {
    'Q1': '一',
    'Q2': '二',
    'Q3': '三',
    'Q4': '四'
}


def doc_info_init(document, year, season, member):
    '''
    Initialize the document information.
    '''
    document.core_properties.author = "John Wang"
    # doc.core_properties.category = "Queclink"
    # doc.core_properties.comments = "Queclink"
    # doc.core_properties.content_status = "Queclink"
    document.core_properties.created = datetime.datetime.now()
    # doc.core_properties.identifier = "Queclink"
    # doc.core_properties.keywords = "Queclink"
    document.core_properties.language = "zh-CN"
    document.core_properties.last_modified_by = "John Wang"
    # doc.core_properties.revision = 1
    # doc.core_properties.subject = "Queclink"
    # doc.core_properties.title = "Queclink"
    # doc.core_properties.version = "Queclink"
    
    replace_list['<<Name>>'] = member.name_cn
    replace_list['<<Year>>'] = year
    replace_list['<<Season>>'] = season_cvt[season]
    
    start_d, end_d = season_date(season)
    replace_list['<<S_Month>>'] = start_d[:2]
    replace_list['<<S_Day>>'] = start_d[2:]
    replace_list['<<E_Month>>'] = end_d[:2]
    replace_list['<<E_Day>>'] = end_d[2:]

def update_doc_info(doc, year, season, member):
    '''
    Update the document information.
    '''
    doc_info_init(doc, year, season, member)
    
    _temp_list = []
    for i,j in replace_list.items():
        _temp_list.append(tuple([i, j]))
    print("_temp_list: " + str(_temp_list))

    table_count = 0
    for table in doc.tables:
        table_count += 1
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)
                for tl in _temp_list:
                    if tl[0] in cell.text:
                        cell.text = cell.text.replace(tl[0], tl[1])
                        print(f"Table Cell1 Replace {tl[0]} with {tl[1]}")
                for para in cell.paragraphs:
                    if table_count <= 2:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for para in doc.paragraphs:
        for tl in _temp_list:
            if tl[0] in para.text:
                para.text = para.text.replace(tl[0], tl[1])
                print(f"Replace {tl[0]} with {tl[1]}")


def folder_init(p_folder):
    '''
    Initialize the folder.
    if not exist, create it.
    if exist, clean the files in it.
    '''
    if not os.path.exists(p_folder):
        os.makedirs(p_folder)
    else:
        for root, dirs, files in os.walk(p_folder):
            for file in files:
                os.remove(os.path.join(root, file))

def kpi_interview_form_generator(docx_template, year, season, rel_path, members):
    '''
    Generate the KPI Interview Form.
    2024年三季度绩效考核面谈表-曹政
    '''
    file_name_prefix = "{}年{}季度绩效考核面谈表-".format(year, season_cvt[season])
    _rpath = os.path.join(rel_path, "interview_form")
    folder_init(_rpath)
    sys_path = os.path.join(_rpath, "sys")
    folder_init(sys_path)
    app_path = os.path.join(_rpath, "app")
    folder_init(app_path)
    
    for member in members:
        _fname = file_name_prefix + member.name_cn + ".docx"
        if member.group == member.GROUP_SYSTEM:
            _rp = os.path.join(sys_path, _fname)
        else:
            _rp = os.path.join(app_path, _fname)
        print(_rp)
        doc = Document(docx_template)
        try:
            update_doc_info(doc, year, season, member)
            doc.save(_rp)
        finally:
            if doc:
                doc.save(_rp)

def build_docx(members, year, season,  rel_path, docx_template = "kpi_interview_form_template.docx"):
    '''     
    Build the KPI Interview Form.
    '''
    kpi_interview_form_generator(docx_template, year, season, rel_path, members)  # Generate the KPI Interview Form.

    
